# gerador_provas_completo.py - Versão final e corrigida
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import requests
from io import BytesIO
import tempfile
from PIL import Image

class GeradorProvasApp:
    def __init__(self, root):
        self.root = root
        self.root.title("📝 Gerador de Provas - Dr. Edesio Martins v2.0")
        self.root.geometry("1100x850")
        self.root.configure(bg='#f0f0f0')

        style = ttk.Style()
        style.theme_use('clam')

        self.questoes = []
        self.logo_path = None
        self.disciplinas = []
        
        self.criar_interface()

    def carregar_logo(self):
        path = filedialog.askopenfilename(
            title="Selecionar Logo",
            filetypes=[("Imagens", "*.png *.jpg *.jpeg *.gif *.bmp")]
        )
        if path:
            self.logo_path = path
            nome_arquivo = os.path.basename(path)
            self.label_logo.config(text=f"✅ {nome_arquivo}", fg="green")
            messagebox.showinfo("✅ Sucesso", f"Logo carregado: {nome_arquivo}")

    def adicionar_disciplina(self):
        disciplina_nome = self.disciplina_entry.get().strip()
        if disciplina_nome and disciplina_nome not in self.disciplinas:
            self.disciplinas.append(disciplina_nome)
            self.disciplinas_listbox.insert(tk.END, disciplina_nome)
            self.disciplina_entry.delete(0, tk.END)

    def remover_disciplina(self):
        try:
            index = self.disciplinas_listbox.curselection()[0]
            self.disciplinas_listbox.delete(index)
            self.disciplinas.pop(index)
        except IndexError:
            pass

    def adicionar_questao(self):
        self._abrir_janela_questao()

    def editar_questao(self):
        try:
            index_selecionado = self.questoes_listbox.curselection()[0]
            if not self.questoes_listbox.get(index_selecionado).startswith("  "):
                messagebox.showwarning("Aviso", "Por favor, selecione uma questão para editar, não um separador de disciplina.")
                return

            num_questao_visual = 0
            questao_para_editar = None
            
            for i, item in enumerate(self.questoes_listbox.get(0, tk.END)):
                if item.startswith("  "):
                    num_questao_visual += 1
                    if i == index_selecionado:
                        questao_para_editar = self.questoes[num_questao_visual - 1]
                        self._abrir_janela_questao(questao_para_editar, num_questao_visual - 1)
                        return
            
            messagebox.showwarning("Aviso", "Questão não encontrada. Tente novamente.")

        except IndexError:
            messagebox.showwarning("Aviso", "Por favor, selecione uma questão para editar.")

    def _abrir_janela_questao(self, questao_a_editar=None, index_a_editar=None):
        janela_questao = tk.Toplevel(self.root)
        janela_questao.title("Adicionar/Editar Questão")
        janela_questao.geometry("700x600")
        
        # Torna a janela um diálogo modal
        janela_questao.grab_set()
        janela_questao.transient(self.root)

        if not self.disciplinas:
            messagebox.showwarning("Aviso", "Por favor, adicione ao menos uma disciplina na aba 'Dados da Prova' primeiro.")
            janela_questao.destroy()
            return
            
        frame_selecao_disciplina = ttk.LabelFrame(janela_questao, text="Selecione a Disciplina", padding=10)
        frame_selecao_disciplina.pack(fill='x', padx=10, pady=5)
        disciplina_selecionada = tk.StringVar()
        combo_disciplina = ttk.Combobox(frame_selecao_disciplina, textvariable=disciplina_selecionada, values=self.disciplinas, state="readonly")
        combo_disciplina.pack(fill='x')
        if self.disciplinas:
            combo_disciplina.set(self.disciplinas[0])

        frame_conteudo_questao = ttk.Frame(janela_questao, padding=10)
        frame_conteudo_questao.pack(fill='both', expand=True, padx=10, pady=10)

        tk.Label(frame_conteudo_questao, text="Enunciado:", font=('Arial', 10, 'bold')).pack(anchor="w")
        enunciado = tk.Text(frame_conteudo_questao, height=4, wrap=tk.WORD, font=('Arial', 10))
        enunciado.pack(fill="x", pady=5)

        tk.Label(frame_conteudo_questao, text="Alternativas:", font=('Arial', 10, 'bold')).pack(anchor="w", pady=(10,5))
        
        alternativas = {}
        correta = tk.StringVar()
        correta.set("A")

        for i, alt in enumerate(["A", "B", "C", "D"]):
            row = tk.Frame(frame_conteudo_questao)
            row.pack(fill="x", pady=2)
            
            rb = tk.Radiobutton(row, text=f"{alt})", variable=correta, value=alt, 
                               font=('Arial', 10, 'bold'), fg="blue")
            rb.pack(side="left", padx=(0, 10))
            
            entry = tk.Entry(row, width=80, font=('Arial', 10))
            entry.pack(side="left", fill="x", expand=True)
            alternativas[alt] = entry

        frame_img = ttk.LabelFrame(frame_conteudo_questao, text="🖼️ Imagem (Opcional)", padding=10)
        frame_img.pack(fill="x", pady=10)

        img_path = tk.StringVar()
        img_url = tk.StringVar()

        def carregar_img():
            path = filedialog.askopenfilename(title="Selecionar Imagem", filetypes=[("Imagens", "*.png *.jpg *.jpeg *.gif *.bmp")])
            if path:
                img_path.set(path)
                nome = os.path.basename(path)
                btn_img.config(text=f"✅ {nome}")

        btn_img = ttk.Button(frame_img, text="📂 Carregar do Computador", command=carregar_img)
        btn_img.pack(side="left", padx=5)

        tk.Label(frame_img, text="ou URL:", font=('Arial', 10)).pack(side="left", padx=(20, 5))
        ttk.Entry(frame_img, textvariable=img_url, width=50).pack(side="left", fill="x", expand=True, padx=5)

        if questao_a_editar:
            janela_questao.title("Editar Questão")
            disciplina_selecionada.set(questao_a_editar['disciplina'])
            enunciado.insert('1.0', questao_a_editar['enunciado'])
            correta.set(questao_a_editar['correta'])
            for alt, texto in questao_a_editar['alternativas'].items():
                if texto:
                    alternativas[alt].insert(0, texto)
            img_path.set(questao_a_editar['img_path'])
            img_url.set(questao_a_editar['img_url'])

        def salvar_questao():
            if not enunciado.get("1.0", "end-1c").strip():
                messagebox.showerror("Erro", "O enunciado não pode ser vazio.")
                return
            
            disciplina_escolhida = disciplina_selecionada.get()
            if not disciplina_escolhida:
                messagebox.showerror("Erro", "Selecione uma disciplina.")
                return

            nova_questao = {
                "disciplina": disciplina_escolhida,
                "enunciado": enunciado.get("1.0", "end-1c"),
                "alternativas": {alt: entry.get() for alt, entry in alternativas.items()},
                "correta": correta.get(),
                "img_path": img_path.get(),
                "img_url": img_url.get()
            }

            if questao_a_editar:
                self.questoes[index_a_editar] = nova_questao
            else:
                self.questoes.append(nova_questao)

            self.atualizar_listagem_questoes()
            self.atualizar_contador()
            
            if not questao_a_editar:
                enunciado.delete('1.0', 'end')
                for entry in alternativas.values():
                    entry.delete(0, 'end')
                correta.set("A")
                img_path.set("")
                img_url.set("")
            
        def concluir():
            janela_questao.destroy()

        frame_botoes = ttk.Frame(janela_questao)
        frame_botoes.pack(pady=10)

        ttk.Button(frame_botoes, text="Salvar Questão", command=salvar_questao).pack(side="left", padx=5)
        ttk.Button(frame_botoes, text="Concluir", command=concluir).pack(side="left", padx=5)

    def criar_interface(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)

        self.aba_dados = ttk.Frame(notebook)
        self.aba_dados.pack(fill='both', expand=True)
        notebook.add(self.aba_dados, text="📋 Dados da Prova")
        self.criar_aba_dados()

        self.aba_questoes = ttk.Frame(notebook)
        self.aba_questoes.pack(fill='both', expand=True)
        notebook.add(self.aba_questoes, text="❓ Questões")
        self.criar_aba_questoes()

        self.aba_config = ttk.Frame(notebook)
        self.aba_config.pack(fill='both', expand=True)
        notebook.add(self.aba_config, text="⚙️ Configurações")
        self.criar_aba_configuracoes()

    def criar_aba_dados(self):
        canvas = tk.Canvas(self.aba_dados, bg='white')
        scrollbar = ttk.Scrollbar(self.aba_dados, orient="vertical", command=canvas.yview)
        self.scrollable_frame_dados = ttk.Frame(canvas)

        self.scrollable_frame_dados.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.scrollable_frame_dados, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.bind_all("<MouseWheel>", lambda event: canvas.yview_scroll(int(-1*(event.delta/120)), "units"))
        canvas.bind_all("<Button-4>", lambda event: canvas.yview_scroll(-1, "units"))
        canvas.bind_all("<Button-5>", lambda event: canvas.yview_scroll(1, "units"))

        frame_info = ttk.LabelFrame(self.scrollable_frame_dados, text="🏫 Informações da Prova", padding=15)
        frame_info.pack(fill="x", padx=10, pady=10)

        self.instituicao = tk.StringVar(value="")
        self.disciplina = tk.StringVar()
        self.periodo = tk.StringVar()
        self.professor = tk.StringVar(value="")
        self.data = tk.StringVar(value="dd/mm/aaaa")

        tk.Label(frame_info, text="Instituição:", font=('Arial', 10, 'bold')).grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(frame_info, textvariable=self.instituicao, width=50, font=('Arial', 10)).grid(row=0, column=1, columnspan=3, padx=5, pady=5, sticky="ew")

        tk.Label(frame_info, text="Disciplina:", font=('Arial', 10, 'bold')).grid(row=1, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(frame_info, textvariable=self.disciplina, width=30, font=('Arial', 10)).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(frame_info, text="Período:", font=('Arial', 10, 'bold')).grid(row=1, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(frame_info, textvariable=self.periodo, width=15, font=('Arial', 10)).grid(row=1, column=3, padx=5, pady=5, sticky="ew")

        tk.Label(frame_info, text="Professor:", font=('Arial', 10, 'bold')).grid(row=2, column=0, sticky="w", padx=5, pady=5)
        ttk.Entry(frame_info, textvariable=self.professor, width=30, font=('Arial', 10)).grid(row=2, column=1, padx=5, pady=5, sticky="ew")

        tk.Label(frame_info, text="Data:", font=('Arial', 10, 'bold')).grid(row=2, column=2, sticky="w", padx=5, pady=5)
        ttk.Entry(frame_info, textvariable=self.data, width=15, font=('Arial', 10)).grid(row=2, column=3, padx=5, pady=5, sticky="ew")
        
        frame_logo = ttk.LabelFrame(self.scrollable_frame_dados, text="🖼️ Logo da Instituição", padding=15)
        frame_logo.pack(fill="x", padx=10, pady=10)

        btn_logo = ttk.Button(frame_logo, text="📂 Carregar Logo", command=self.carregar_logo)
        btn_logo.pack(side="left", padx=5)

        self.label_logo = tk.Label(frame_logo, text="Nenhum logo carregado", fg="gray")
        self.label_logo.pack(side="left", padx=10)
        
        frame_disciplinas = ttk.LabelFrame(self.scrollable_frame_dados, text="📚 Disciplinas da Prova", padding=15)
        frame_disciplinas.pack(fill="x", padx=10, pady=10)
        
        self.disciplina_entry = ttk.Entry(frame_disciplinas, width=40)
        self.disciplina_entry.pack(side="left", padx=5)
        
        ttk.Button(frame_disciplinas, text="➕ Adicionar Disciplina", command=self.adicionar_disciplina).pack(side="left", padx=5)
        
        self.disciplinas_listbox = tk.Listbox(frame_disciplinas, height=5)
        self.disciplinas_listbox.pack(fill="x", pady=5)
        
        ttk.Button(frame_disciplinas, text="🗑️ Remover Selecionada", command=self.remover_disciplina).pack(side="left", padx=5, pady=5)

        frame_instr = ttk.LabelFrame(self.scrollable_frame_dados, text="📝 Instruções Gerais", padding=15)
        frame_instr.pack(fill="x", padx=10, pady=10)

        self.instrucoes = tk.Text(frame_instr, height=10, wrap=tk.WORD, font=('Arial', 10))
        scroll_instr = ttk.Scrollbar(frame_instr, orient="vertical", command=self.instrucoes.yview)
        self.instrucoes.configure(yscrollcommand=scroll_instr.set)

        self.instrucoes.pack(side="left", fill="both", expand=True)
        scroll_instr.pack(side="right", fill="y")

        instrucoes_padrao = """1. Leia atentamente cada questão da prova antes de começar a respondê-la. Comece respondendo à questão cuja resposta esteja clara na memória. Desta forma, você aproveitará melhor o tempo de prova.
2. A avaliação deverá ser feita, individualmente, a caneta azul ou preta, somente será aceito gabarito preenchido nestas cores.
3. Para cada questão existe apenas uma alternativa que a responde acertadamente. Para a marcação da alternativa escolhida pinte a FOLHA DE RESPOSTAS completamente no campo correspondente.
4. Será invalidada a questão em que houver mais de uma marcação, marcação rasurada ou emendada, ou não houver marcação. Não serão aceitas rasuras, as respostas com rasuras não serão computadas.
5. A duração da prova é de acordo com o horário disponibilizado na programação da semana de provas, já incluído o tempo destinado ao preenchimento da FOLHA DE RESPOSTAS. Não haverá prorrogação de tempo!
6. Terminada a prova, você deverá, OBRIGATORIAMENTE, entregar a FOLHA DE RESPOSTAS ao fiscal responsável de sua sala.
7. Apenas será disponibilizada uma FOLHA DE RESPOSTAS (atenção ao marcar suas respostas).
8. É proibido portar qualquer instrumento eletrônico, incluindo celulares e relógios, durante a aplicação das provas. Os mesmos devem permanecer desligados fora da sala de aula ou em local orientado pelo professor. Qualquer tipo de atitude, por parte do aluno, na tentativa de “cola”, uso de celular ou caso o aluno seja flagrado desrespeitando as normas ditadas pelo professor responsável, será impedido de continuar a prova e terá nota ZERO.
9. Diante de qualquer dúvida, você deve comunicar-se com o professor responsável pela aplicação da prova.
10. O TEMPO DE DURAÇÃO mínima da avaliação é de 30 minutos, contados a partir do início da avaliação."""
        self.instrucoes.insert("1.0", instrucoes_padrao)

        frame_cabecalho = ttk.LabelFrame(self.scrollable_frame_dados, text="📄 Notas no Cabeçalho", padding=15)
        frame_cabecalho.pack(fill="x", padx=10, pady=10)

        tk.Label(frame_cabecalho, text="Texto para cabeçalho (a partir da 2ª página):", font=('Arial', 10, 'bold')).pack(anchor="w")
        self.notas_cabecalho = tk.StringVar(value="")
        ttk.Entry(frame_cabecalho, textvariable=self.notas_cabecalho, width=80, font=('Arial', 10)).pack(fill="x", pady=5)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def criar_aba_questoes(self):
        frame_controles = ttk.Frame(self.aba_questoes)
        frame_controles.pack(fill="x", padx=10, pady=10)

        ttk.Button(frame_controles, text="➕ Adicionar Questões", command=self.adicionar_questao).pack(side="left", padx=5)
        ttk.Button(frame_controles, text="✏️ Editar Questão", command=self.editar_questao).pack(side="left", padx=5)
        ttk.Button(frame_controles, text="📊 Ver Gabarito", command=self.mostrar_gabarito).pack(side="left", padx=5)
        ttk.Button(frame_controles, text="🗑️ Limpar Tudo", command=self.limpar_questoes).pack(side="left", padx=5)

        self.label_contador = tk.Label(frame_controles, text="Questões: 0", font=('Arial', 10, 'bold'), fg="blue")
        self.label_contador.pack(side="right", padx=10)

        canvas = tk.Canvas(self.aba_questoes, bg='white')
        self.questoes_listbox = tk.Listbox(canvas)
        self.questoes_listbox.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.questoes_listbox.bind("<Double-1>", lambda event: self.editar_questao())
        
        self.questoes_listbox_scrollbar = ttk.Scrollbar(canvas, orient="vertical", command=self.questoes_listbox.yview)
        self.questoes_listbox.configure(yscrollcommand=self.questoes_listbox_scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        self.questoes_listbox_scrollbar.pack(side="right", fill="y")
        
        self.atualizar_listagem_questoes()

    def criar_aba_configuracoes(self):
        frame_geracao = ttk.LabelFrame(self.aba_config, text="📄 Gerar Documentos", padding=20)
        frame_geracao.pack(fill="x", padx=20, pady=20)

        ttk.Button(frame_geracao, text="📝 Gerar Prova em Word (.docx)",
                   command=self.gerar_word, style="Accent.TButton").pack(pady=5, fill="x")

        frame_opcoes = ttk.LabelFrame(self.aba_config, text="⚙️ Opções de Formatação", padding=20)
        frame_opcoes.pack(fill="x", padx=20, pady=20)

        self.duas_colunas = tk.BooleanVar(value=True)
        self.incluir_gabarito = tk.BooleanVar(value=True)
        self.numerar_paginas = tk.BooleanVar(value=True)

        ttk.Checkbutton(frame_opcoes, text="Questões em duas colunas (a partir da 2ª página)",
                        variable=self.duas_colunas).pack(anchor="w", pady=2)
        ttk.Checkbutton(frame_opcoes, text="Incluir gabarito",
                        variable=self.incluir_gabarito).pack(anchor="w", pady=2)

        frame_info = ttk.LabelFrame(self.aba_config, text="ℹ️ Sobre", padding=20)
        frame_info.pack(fill="x", padx=20, pady=20)

        info_text = """Gerador de Provas v2.0
Desenvolvido por Dr. Edesio Martins
• Suporte a imagens locais e URLs
• Geração de Word formatado
• Interface moderna e intuitiva"""

        tk.Label(frame_info, text=info_text, justify="left", font=('Arial', 10)).pack(anchor="w")

    def atualizar_listagem_questoes(self):
        self.questoes_listbox.delete(0, tk.END)
        
        questoes_por_disciplina = {}
        for q in self.questoes:
            disciplina = q["disciplina"]
            if disciplina not in questoes_por_disciplina:
                questoes_por_disciplina[disciplina] = []
            questoes_por_disciplina[disciplina].append(q)

        num_questao_geral = 1
        for disciplina in self.disciplinas:
            self.questoes_listbox.insert(tk.END, f"--- {disciplina.upper()} ---")
            if disciplina in questoes_por_disciplina:
                for q in questoes_por_disciplina[disciplina]:
                    self.questoes_listbox.insert(tk.END, f"  Questão {num_questao_geral}: {q['enunciado'][:50]}...")
                    num_questao_geral += 1
            
    def remover_questao(self, questao_obj):
        if messagebox.askyesno("⚠️ Confirmar", "Deseja realmente remover esta questão?"):
            self.questoes.remove(questao_obj)
            self.atualizar_listagem_questoes()
            self.atualizar_contador()

    def atualizar_contador(self):
        total = len(self.questoes)
        self.label_contador.config(text=f"Questões: {total}")

    def limpar_questoes(self):
        if messagebox.askyesno("⚠️ Confirmar", "Deseja realmente remover todas as questões?"):
            self.questoes.clear()
            self.atualizar_listagem_questoes()
            self.atualizar_contador()

    def mostrar_gabarito(self):
        if not self.questoes:
            messagebox.showwarning("⚠️ Aviso", "Adicione questões primeiro!")
            return

        gabarito = []
        for i, q in enumerate(self.questoes, 1):
            resposta = q["correta"] or "Não marcada"
            gabarito.append(f"Questão {i}: {resposta}")

        janela_gabarito = tk.Toplevel(self.root)
        janela_gabarito.title("📊 Gabarito")
        janela_gabarito.geometry("300x400")

        text_widget = tk.Text(janela_gabarito, font=('Arial', 12))
        text_widget.pack(fill="both", expand=True, padx=10, pady=10)

        text_widget.insert("1.0", "\n".join(gabarito))
        text_widget.config(state="disabled")

    def validar_dados(self):
        if not self.questoes:
            messagebox.showerror("❌ Erro", "Adicione ao menos uma questão!")
            return False
        
        if not self.disciplina.get().strip() and not self.disciplinas:
            messagebox.showerror("❌ Erro", "Preencha o nome da disciplina principal ou adicione as disciplinas!")
            return False

        return True

    def gerar_word(self):
        if not self.validar_dados():
            return

        try:
            arquivo = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Documentos Word", "*.docx")],
                title="Salvar Prova em Word"
            )

            if not arquivo:
                return

            self._criar_documento_word(arquivo)
            messagebox.showinfo("✅ Sucesso", f"Prova em Word gerada:\n{arquivo}")

        except Exception as e:
            messagebox.showerror("❌ Erro", f"Erro ao gerar Word:\n{str(e)}")

    def _criar_documento_word(self, arquivo):
        doc = Document()
        
        section_page1 = doc.sections[0]
        section_page1.left_margin = Cm(1.5)
        section_page1.right_margin = Cm(1.5)
        
        if self.logo_path and os.path.exists(self.logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
            run.add_picture(self.logo_path, width=Inches(2))

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = titulo.add_run(self.instituicao.get())
        run_titulo.font.size = Inches(0.2)
        run_titulo.bold = True

        doc.add_paragraph(f"Disciplina: {self.disciplina.get()}", style='Normal')
        doc.add_paragraph(f"Período: {self.periodo.get()}", style='Normal')
        doc.add_paragraph(f"Professor: {self.professor.get()}", style='Normal')
        doc.add_paragraph(f"Data: {self.data.get()}", style='Normal')
        
        p_aluno = doc.add_paragraph()
        p_aluno.add_run("Aluno: ____________________________________________________________ ")
        p_aluno.add_run("Matrícula: __________________________________")

        doc.add_paragraph()
        instr_titulo = doc.add_paragraph()
        instr_titulo.add_run("INSTRUÇÕES:").bold = True
        
        instrucoes_texto = self.instrucoes.get("1.0", "end-1c")
        p_instrucoes = doc.add_paragraph(instrucoes_texto)
        p_instrucoes.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()
        section_page2_onward = doc.add_section(WD_SECTION.CONTINUOUS)
        section_page2_onward.left_margin = Cm(1.5)
        section_page2_onward.right_margin = Cm(1.5)

        header = section_page2_onward.header
        header.is_linked_to_previous = False
        
        if self.notas_cabecalho.get():
            header_para = header.paragraphs[0]
            header_para.text = self.notas_cabecalho.get()

        if self.duas_colunas.get():
            sectPr = section_page2_onward._sectPr
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:sep'), '1')
            cols.set(qn('w:space'), '720')

        gabarito = []
        questao_num = 1
        
        questoes_por_disciplina = {}
        for q in self.questoes:
            disciplina = q["disciplina"]
            if disciplina not in questoes_por_disciplina:
                questoes_por_disciplina[disciplina] = []
            questoes_por_disciplina[disciplina].append(q)
        
        for disciplina in self.disciplinas:
            if disciplina in questoes_por_disciplina:
                if disciplina:
                    p_disc = doc.add_paragraph()
                    p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_disc = p_disc.add_run(f"DISCIPLINA: {disciplina.upper()}")
                    run_disc.bold = True
                    doc.add_paragraph()
                
                for q in questoes_por_disciplina[disciplina]:
                    questao_para = doc.add_paragraph()
                    run_num = questao_para.add_run(f"Questão {questao_num:02d}. ")
                    run_num.bold = True
                    run_num.font.size = Inches(0.15)
            
                    enunciado_texto = q["enunciado"]
                    p_enunciado = doc.add_paragraph(enunciado_texto)
                    run_enunciado = p_enunciado.runs[0] if p_enunciado.runs else p_enunciado.add_run()
                    run_enunciado.bold = True
                    
                    p_enunciado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_enunciado.paragraph_format.line_spacing = 1.0
                    p_enunciado.paragraph_format.first_line_indent = Cm(0.63)
            
                    self._adicionar_imagem_word(doc, q)
            
                    for alt in ["A", "B", "C", "D"]:
                        texto_alt = q["alternativas"].get(alt, "").strip()
                        if texto_alt:
                            p_alt = doc.add_paragraph(f"{alt}) {texto_alt}")
                            p_alt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_alt.paragraph_format.line_spacing = 1.0
                            p_alt.paragraph_format.first_line_indent = Cm(0.63)
            
                    doc.add_paragraph()
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.paragraph_format.space_after = Cm(0.63)
                    
                    gabarito.append((questao_num, q["correta"]))
                    questao_num += 1

        if self.incluir_gabarito.get():
            doc.add_page_break()
            section_gabarito = doc.add_section(WD_SECTION.CONTINUOUS)
            section_gabarito.left_margin = Cm(1.5)
            section_gabarito.right_margin = Cm(1.5)

            gab_titulo = doc.add_paragraph()
            gab_run = gab_titulo.add_run("GABARITO")
            gab_run.bold = True
            gab_run.font.size = Inches(0.2)
            gab_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Questão'
            hdr_cells[1].text = 'Resposta'

            for questao_num, resposta in gabarito:
                row_cells = table.add_row().cells
                row_cells[0].text = str(questao_num)
                row_cells[1].text = resposta if resposta else "—"
        
        doc.save(arquivo)

    def _adicionar_imagem_word(self, doc, questao):
        try:
            img_path = questao.get("img_path", "")
            img_url = questao.get("img_url", "")

            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(3))
            elif img_url:
                response = requests.get(img_url, timeout=10)
                if response.status_code == 200:
                    img_bytes = BytesIO(response.content)
                    doc.add_picture(img_bytes, width=Inches(3))
                else:
                    doc.add_paragraph("[Erro: Não foi possível carregar a imagem da URL]")
        except Exception as e:
            doc.add_paragraph(f"[Erro ao carregar imagem: {str(e)}]")


def main():
    root = tk.Tk()
    try:
        if os.path.exists('icone.ico'):
            root.iconbitmap('icone.ico')
    except:
        pass

    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')

    app = GeradorProvasApp(root)
    app.atualizar_listagem_questoes()
    root.mainloop()

if __name__ == "__main__":
    main()